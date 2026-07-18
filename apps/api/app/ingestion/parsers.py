"""Extract plain text from uploaded documents for chunking/embedding."""

from __future__ import annotations

import json
import re
from pathlib import Path

# Plain text / markup (read as UTF-8)
TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".log",
    ".xml",
    ".html",
    ".htm",
    ".yml",
    ".yaml",
    ".toml",
    ".ini",
    ".cfg",
    ".css",
    ".js",
    ".ts",
    ".py",
    ".sh",
}

# Binary-ish office / PDF
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}  # not legacy .doc (binary OLE)

SUPPORTED_SUFFIXES = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES

# For file input accept= attribute
ACCEPT_EXTENSIONS = ",".join(sorted(SUPPORTED_SUFFIXES))


class ParseError(Exception):
    """Raised when a document cannot be parsed."""


def is_supported_filename(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_SUFFIXES


def supported_types_message() -> str:
    return (
        "Supported: PDF, DOCX, and text-like files "
        f"({', '.join(sorted(TEXT_SUFFIXES))}). "
        "Legacy .doc is not supported — save as .docx."
    )


def parse_file(path: Path) -> str:
    """Read a local file and return plain text for RAG."""
    if not path.is_file():
        raise ParseError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ParseError(
            f"Unsupported type '{suffix}'. {supported_types_message()}"
        )

    if suffix in PDF_SUFFIXES:
        text = _parse_pdf(path)
    elif suffix in DOCX_SUFFIXES:
        text = _parse_docx(path)
    elif suffix in {".html", ".htm"}:
        text = _parse_html(path)
    elif suffix == ".json":
        text = _parse_json(path)
    else:
        text = _parse_text(path)

    text = _normalize_whitespace(text)
    if not text.strip():
        raise ParseError(
            "No extractable text found. "
            "Scanned PDFs without OCR, empty files, or image-only docs will fail."
        )
    return text


def _parse_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ParseError("PDF support requires pypdf. Install API deps.") from e

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ParseError(f"Could not open PDF: {e}") from e

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = page_text.strip()
        if page_text:
            parts.append(f"--- Page {i + 1} ---\n{page_text}")

    return "\n\n".join(parts)


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ParseError("DOCX support requires python-docx.") from e

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ParseError(f"Could not open DOCX: {e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # Tables → plain text rows
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    return "\n\n".join(parts)


def strip_html_text(raw: str) -> str:
    # Lightweight strip tags (no BeautifulSoup dep)
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", raw)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"&#\d+;", " ", text)
    return text


def _parse_html(path: Path) -> str:
    return strip_html_text(_parse_text(path))


def _parse_json(path: Path) -> str:
    raw = _parse_text(path)
    try:
        data = json.loads(raw)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return raw


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
